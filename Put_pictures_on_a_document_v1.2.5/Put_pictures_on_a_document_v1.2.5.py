from ast import Pass
import os
from docx import Document
from docx.shared import Inches, Cm
import time, sys

document = Document()

#Info printing
print("This program can put 4 pictures on a docx document!\n\n")

#Path processing
pgm_path = sys.argv[0]
pgm_path_list = []
pgm_path_list = pgm_path.split('\\')
tmp = pgm_path_list.pop()
tmp_len = len(tmp)
pgm_path = pgm_path[0:(len(pgm_path)-tmp_len)]

#Open image folder
os.startfile(pgm_path+'images')
qr_fname = str(input('Drag the picture here:'))

print("Start creating document...")

#Page settings
section = document.sections[0]
section.page_height=Cm(29.7)
section.page_width=Cm(21)
section.top_margin=Cm(0.6)
section.bottom_margin=Cm(0.6)
section.left_margin=Cm(0.6)
section.right_margin=Cm(0.6)

#Create table to hold pictures
table = document.add_table(rows=2,cols=2)
table.style = 'Table Grid'

#Identify pictures
try:
    if '"' in qr_fname:
        tmplist = qr_fname.split('"')
        qr_str = tmplist[1]
    elif '&' in qr_fname:
        tmplist = qr_fname.split("'")
        qr_str = tmplist[-2]
    else:
        qr_str=qr_fname
except:
    print("Something went wrong...")

#Fill the table with pictures
try:
    for i in range(0,2):
        for j in range(0,2):
            cell = table.cell(i,j)
            ph = cell.paragraphs[0]
            run = ph.add_run('')
            run.add_picture(qr_str,width=Cm(8.45))
except FileNotFoundError:
    print("We couldn't find the picture in the folder: 'images'")
    time.sleep(1)
    sys.exit()
except OSError:
    print("Format error!")
    time.sleep(1)
    sys.exit()

#Document finished
print("The document is created successfully!")
os.startfile(pgm_path+'docs')
time.sleep(1)

#Save document
time_info = str(time.strftime("%y%m%d",time.localtime()))
out_str = 'docs\\' + time_info + '.docx'
document.save(out_str)
